"""Tests for PDF, DOCX, and pasted resume parsing helpers."""

from io import BytesIO

import pytest
from docx import Document

from src.resume_parser import (
    ResumeParsingError,
    extract_resume_text,
    extract_text_from_docx,
    extract_text_from_pdf,
    validate_resume_text,
)


def test_empty_pdf_is_rejected() -> None:
    with pytest.raises(ResumeParsingError, match="empty"):
        extract_text_from_pdf(BytesIO(b""))


def test_invalid_pdf_is_rejected() -> None:
    with pytest.raises(ResumeParsingError, match="valid PDF"):
        extract_text_from_pdf(BytesIO(b"not a pdf"))


def test_docx_paragraphs_and_tables_are_extracted() -> None:
    document = Document()
    document.add_heading("Jordan Lee", level=1)
    document.add_paragraph("Computer science student with Python project experience.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, SQL, Git, React"
    document.add_paragraph(
        "Built and tested a campus application used by 40 students."
    )
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_docx(buffer.getvalue())

    assert "Jordan Lee" in text
    assert "Skills | Python, SQL, Git, React" in text


def test_unified_parser_rejects_unsupported_extension() -> None:
    with pytest.raises(ResumeParsingError, match="PDF or DOCX"):
        extract_resume_text(b"resume content", "resume.txt")


def test_pasted_resume_requires_enough_content() -> None:
    with pytest.raises(ResumeParsingError, match="too little"):
        validate_resume_text("Python student")
