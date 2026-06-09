"""Validate and extract resume text from PDF, DOCX, or pasted text."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import BinaryIO
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024
MAX_PAGES = 10
MAX_TEXT_CHARACTERS = 60_000
SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeParsingError(ValueError):
    """Raised when resume content cannot be safely parsed."""


def extract_resume_text(file: BinaryIO | bytes, filename: str) -> str:
    """Extract resume text based on the uploaded filename."""

    extension = Path(filename).suffix.casefold()
    if extension not in SUPPORTED_EXTENSIONS:
        raise ResumeParsingError("Upload a PDF or DOCX resume.")
    if extension == ".pdf":
        return extract_text_from_pdf(file)
    return extract_text_from_docx(file)


def extract_text_from_pdf(file: BinaryIO | bytes) -> str:
    """Extract readable text from a PDF resume."""

    raw_bytes = _read_and_validate_file(file)
    try:
        reader = PdfReader(BytesIO(raw_bytes))
    except (PdfReadError, ValueError, TypeError) as exc:
        raise ResumeParsingError("The uploaded file is not a valid PDF.") from exc

    if reader.is_encrypted:
        try:
            decrypt_result = reader.decrypt("")
        except Exception as exc:
            raise ResumeParsingError("Password-protected PDFs are not supported.") from exc
        if not decrypt_result:
            raise ResumeParsingError("Password-protected PDFs are not supported.")

    if len(reader.pages) > MAX_PAGES:
        raise ResumeParsingError("Please upload a resume with 10 pages or fewer.")

    try:
        page_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise ResumeParsingError("The PDF text could not be extracted.") from exc
    return validate_resume_text("\n".join(page_text), source="PDF")


def extract_text_from_docx(file: BinaryIO | bytes) -> str:
    """Extract paragraphs and table content from a DOCX resume."""

    raw_bytes = _read_and_validate_file(file)
    try:
        document = Document(BytesIO(raw_bytes))
    except (PackageNotFoundError, BadZipFile, ValueError, KeyError, OSError) as exc:
        raise ResumeParsingError("The uploaded file is not a valid DOCX document.") from exc

    content = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            content.append(" | ".join(cell.text for cell in row.cells))
    return validate_resume_text("\n".join(content), source="DOCX")


def validate_resume_text(text: str, source: str = "pasted text") -> str:
    """Clean and validate extracted or pasted resume text."""

    cleaned = _clean_text(text)
    if not cleaned:
        raise ResumeParsingError(f"No readable text was found in the {source}.")
    if len(cleaned) > MAX_TEXT_CHARACTERS:
        raise ResumeParsingError(
            "The resume text is too long. Keep it under 60,000 characters."
        )
    if len(cleaned.split()) < 20:
        raise ResumeParsingError(
            "The resume contains too little readable text for a useful analysis."
        )
    return cleaned


def _read_and_validate_file(file: BinaryIO | bytes) -> bytes:
    if isinstance(file, bytes):
        raw_bytes = file
    else:
        if hasattr(file, "seek"):
            file.seek(0)
        raw_bytes = file.read()
    if not raw_bytes:
        raise ResumeParsingError("The uploaded resume is empty.")
    if len(raw_bytes) > MAX_FILE_SIZE_BYTES:
        raise ResumeParsingError("The resume must be smaller than 5 MB.")
    return raw_bytes


def _clean_text(text: str) -> str:
    """Remove repeated whitespace while preserving useful line breaks."""

    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()
